"""Secure file storage and document text extraction for Brainstorm."""

from dataclasses import dataclass
import asyncio
import mimetypes
from pathlib import Path
import re
import uuid

import aiofiles
from docx import Document
from fastapi import UploadFile

from app.config import (
    BRAINSTORM_DOCUMENT_DIR,
    BRAINSTORM_IMAGE_DIR,
    UPLOAD_ROOT,
    settings,
)


class FileStorageError(Exception):
    """Raised when an uploaded file is invalid or cannot be stored."""


@dataclass(slots=True)
class StoredUpload:
    original_filename: str
    stored_filename: str
    file_type: str
    mime_type: str
    file_size: int
    storage_path: str


class FileService:
    """Validate, store, read, and safely delete Brainstorm uploads."""

    IMAGE_TYPES = {"png", "jpg", "jpeg", "webp"}
    DOCUMENT_TYPES = {"pdf", "txt", "md", "docx"}
    MIME_TYPES = {
        "pdf": {"application/pdf"},
        "png": {"image/png"},
        "jpg": {"image/jpeg"},
        "jpeg": {"image/jpeg"},
        "webp": {"image/webp"},
        "txt": {"text/plain"},
        "md": {"text/markdown", "text/plain", "text/x-markdown"},
        "docx": {"application/vnd.openxmlformats-officedocument.wordprocessingml.document"},
    }

    def __init__(self) -> None:
        self.allowed_file_types = {file_type.lower() for file_type in settings.ALLOWED_FILE_TYPES}
        self.max_upload_bytes = settings.MAX_UPLOAD_MB * 1024 * 1024

    async def save_upload(self, upload: UploadFile, image_only: bool = False) -> StoredUpload:
        original_filename, file_type, mime_type = self._validate_upload(upload, image_only=image_only)
        destination_dir = BRAINSTORM_IMAGE_DIR if file_type in self.IMAGE_TYPES else BRAINSTORM_DOCUMENT_DIR
        stored_filename = f"{uuid.uuid4().hex}.{file_type}"
        destination = destination_dir / stored_filename

        total_size = 0
        try:
            async with aiofiles.open(destination, "wb") as output:
                while True:
                    chunk = await upload.read(1024 * 1024)
                    if not chunk:
                        break
                    total_size += len(chunk)
                    if total_size > self.max_upload_bytes:
                        raise FileStorageError(
                            f"File exceeds maximum upload size of {settings.MAX_UPLOAD_MB} MB."
                        )
                    await output.write(chunk)
        except Exception:
            if destination.exists():
                destination.unlink(missing_ok=True)
            raise

        if total_size <= 0:
            destination.unlink(missing_ok=True)
            raise FileStorageError("Uploaded file is empty.")

        return StoredUpload(
            original_filename=original_filename,
            stored_filename=stored_filename,
            file_type=file_type,
            mime_type=mime_type,
            file_size=total_size,
            storage_path=str(destination.resolve()),
        )

    async def extract_plain_document_text(self, path: str | Path, file_type: str) -> str:
        file_path = Path(path)
        if file_type in {"txt", "md"}:
            return await asyncio.to_thread(self._read_text_file, file_path)
        if file_type == "docx":
            return await asyncio.to_thread(self._read_docx_file, file_path)
        raise FileStorageError(f"Unsupported document extraction type: {file_type}")

    def delete_file_safely(self, path: str | Path) -> None:
        file_path = Path(path).resolve()
        upload_root = UPLOAD_ROOT.resolve()
        try:
            file_path.relative_to(upload_root)
        except ValueError as exc:
            raise FileStorageError("Refusing to delete a file outside the upload directory.") from exc
        if file_path.exists() and file_path.is_file():
            file_path.unlink()

    def _validate_upload(self, upload: UploadFile, image_only: bool = False) -> tuple[str, str, str]:
        original_filename = self._safe_original_filename(upload.filename)
        suffix = Path(original_filename).suffix.lower().lstrip(".")
        if suffix == "jpeg":
            file_type = "jpeg"
        else:
            file_type = suffix

        if file_type not in self.allowed_file_types:
            raise FileStorageError(f"Unsupported file type: .{suffix or 'unknown'}")
        if image_only and file_type not in self.IMAGE_TYPES:
            raise FileStorageError("Only image files are accepted by this endpoint.")
        if file_type not in self.MIME_TYPES:
            raise FileStorageError(f"File type .{file_type} is not configured.")

        guessed_mime = mimetypes.guess_type(original_filename)[0]
        reported_mime = upload.content_type or guessed_mime or "application/octet-stream"
        mime_type = guessed_mime if reported_mime == "application/octet-stream" and guessed_mime else reported_mime

        if mime_type not in self.MIME_TYPES[file_type]:
            raise FileStorageError(f"Unsupported MIME type for .{file_type}: {mime_type}")

        return original_filename, file_type, mime_type

    def _safe_original_filename(self, filename: str | None) -> str:
        candidate = Path(filename or "upload").name.strip()
        candidate = re.sub(r"[^A-Za-z0-9._ -]+", "_", candidate)
        candidate = re.sub(r"\s+", " ", candidate).strip(" .")
        if not candidate or "." not in candidate:
            raise FileStorageError("Uploaded file must include a supported extension.")
        return candidate[:255]

    def _read_text_file(self, path: Path) -> str:
        data = path.read_bytes()
        for encoding in ("utf-8", "utf-8-sig", "latin-1"):
            try:
                return data.decode(encoding).strip()
            except UnicodeDecodeError:
                continue
        raise FileStorageError("Text file could not be decoded.")

    def _read_docx_file(self, path: Path) -> str:
        try:
            document = Document(path)
        except Exception as exc:
            raise FileStorageError("DOCX file could not be opened.") from exc

        paragraphs = [paragraph.text.strip() for paragraph in document.paragraphs if paragraph.text.strip()]
        table_text: list[str] = []
        for table in document.tables:
            for row in table.rows:
                cells = [cell.text.strip() for cell in row.cells if cell.text.strip()]
                if cells:
                    table_text.append(" | ".join(cells))

        return "\n\n".join(paragraphs + table_text).strip()
